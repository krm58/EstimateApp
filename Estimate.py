# from __future__ import print_function
from kivy.app import App
from kivy.uix.boxlayout import BoxLayout
from kivy.properties import ObjectProperty
from kivy.uix.listview import ListItemButton
from kivy.uix.screenmanager import ScreenManager, Screen
from kivy.lang import Builder
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from kivy.uix.button import Button
from kivy.properties import StringProperty
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from os.path import basename
from email.mime.application import MIMEApplication
import httplib2
import os
import base64
from apiclient import discovery
from oauth2client import tools
from oauth2client.file import Storage
from kivy.uix.popup import Popup
from kivy.uix.label import Label  
import glob

try:
	import argparse
	flags = argparse.ArgumentParser(parents=[tools.argparser]).parse_args()
except ImportError:
	flags = None

class ToolListButton(ListItemButton):
	pass

class PriceListButton(ListItemButton):
	pass

class EmailInterface(BoxLayout, Screen):
	pass

class GoBack(Screen):
	pass

class ScreenManagement(ScreenManager):
	pass

def get_credentials():
		"""Gets valid user credentials from storage.

		If nothing has been stored, or if the stored credentials are invalid,
		the OAuth2 flow is completed to obtain the new credentials.

		Returns:
			Credentials, the obtained credential.
		"""
		SCOPES = 'https://mail.google.com/'

		home_dir = os.path.expanduser('~')
		credential_dir = os.path.join(home_dir, '.credentials')
		if not os.path.exists(credential_dir):
			os.makedirs(credential_dir)
		credential_path = os.path.join(credential_dir,
									   'gmail-python-quickstart.json')
		store = Storage(credential_path)
		credentials = store.get()
		if not credentials or credentials.invalid:
			flow = client.flow_from_clientsecrets(CLIENT_SECRET_FILE, SCOPES)
			flow.user_agent = APPLICATION_NAME
			if flags:
				credentials = tools.run_flow(flow, store, flags)
			else: # Needed only for compatibility with Python 2.6
				credentials = tools.run(flow, store)
			print('Storing credentials to ' + credential_path)
		return credentials

class IntentButton(Button):
	email_recipient = StringProperty()
	email_sender = StringProperty()
	email_subject = StringProperty()
	email_text = StringProperty()
	email_filename = StringProperty()
	email_company = StringProperty()
	email_supply = StringProperty()

	def send_email(self, name, *args):
		"""Send an email message.
		Args:
		service: Authorized Gmail API service instance.
		user_id: User's email address. The special value "me"
		can be used to indicate the authenticated user.
		message: Message to be sent.
		Returns: Sent Message.
		"""	
		credentials = get_credentials()
		http = credentials.authorize(httplib2.Http())
		service = discovery.build('gmail', 'v1', http=http)

		msg = MIMEMultipart()
		msg['From'] = self.email_sender
		recipients = []
		if self.email_recipient:
			recipients.append(self.email_recipient)
		if self.email_company:
			recipients.append(self.email_company)
		if self.email_supply:
			recipients.append(self.email_supply)
		msg['To'] = ", ".join(recipients)
		msg['Subject'] = self.email_subject
		message = self.email_text
		msg.attach(MIMEText(message))
		
		name = name + '.docx'
		with open(name, "rb") as fil:
			part = MIMEApplication(
				fil.read(),
				Name=basename(name)
				)
			part['Content-Disposition'] = 'attachment; filename="%s"' % basename(name)
			msg.attach(part)
		raw = base64.urlsafe_b64encode(msg.as_bytes())
		raw = raw.decode()
		body = {'raw': raw}
		message = (service.users().messages().send(userId="me", body=body).execute())

class ToolDB(BoxLayout, Screen):
	tool_text_input = ObjectProperty()
	tool_price_input = ObjectProperty()
	tool_list = ObjectProperty()
	price_list = ObjectProperty()
	estnumber = StringProperty()

	def __init__(self, **kwargs):
		super(ToolDB, self).__init__(**kwargs)
		self.estcounter = 0
		self.estnumber = "Estimate Option: " + str(self.estcounter+1)
		self.estimatejobsDict = {}
	
	def change_text(self):
		self.estcounter = self.estcounter + 1
		self.estnumber = "Estimate Option: " + str(self.estcounter+1)

	def prepnewlist(self):
		
		self.estimatejobsDict["toollist_"+str(self.estcounter)] = self.tool_list.adapter.data
		self.estimatejobsDict["pricelist_"+str(self.estcounter)] = self.price_list.adapter.data

		self.tool_list.adapter.data = []
		self.price_list.adapter.data = []

	def submit_tool(self):
		# Get the tool's name from textInputs
		tool_name = self.tool_text_input.text
		tool_price = self.tool_price_input.text
		#Add to ListView
		self.tool_list.adapter.data.extend([tool_name])
		self.price_list.adapter.data.extend([tool_price])
		#Reset the ListView
		self.tool_text_input.text = "" 
		self.tool_price_input.text = ""
		self.tool_list._trigger_reset_populate()
		self.price_list._trigger_reset_populate()

	def delete_tool(self):
		# If a list item is selected using tool
		if self.tool_list.adapter.selection:

			#Get the text from the item selected
			selection = self.tool_list.adapter.selection[0].text
			ind = self.tool_list.adapter.data.index(selection)
			selection_price = self.price_list.adapter.data[ind]
			#Remove the matching item
			self.tool_list.adapter.data.remove(selection)
			self.price_list.adapter.data.remove(selection_price)
			#Reset the List View
			self.tool_list._trigger_reset_populate()
			self.price_list._trigger_reset_populate()

		if self.price_list.adapter.selection: #if list item is selected using price

			#Get the text from the item selected
			selection_price = self.price_list.adapter.selection[0].text
			ind = self.price_list.adapter.data.index(selection_price)
			selection_tool = self.tool_list.adapter.data[ind]
			#Remove the matching item
			self.tool_list.adapter.data.remove(selection_tool)
			self.price_list.adapter.data.remove(selection_price)
			#Reset the List View
			self.tool_list._trigger_reset_populate()
			self.price_list._trigger_reset_populate()

	def save(self,name):
		document = Document()
		my_image = document.add_picture('companylogo.png', width=Inches(1.0))
		last_paragraph = document.paragraphs[-1]
		last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
		for i in range(1,self.estcounter+1):
			paragraph1 = document.add_paragraph()
			if self.estcounter == 1: #if there is only one estimate option
				paragraph1.add_run('Materials List').bold = True
			else:
				headername = "Materials List--Option " + str(i)
				paragraph1.add_run(headername).bold = True
			paragraph1.alignment = WD_ALIGN_PARAGRAPH.CENTER
			#Make a table that includes both the item and the price
			table = document.add_table(rows=1, cols=2)
			table.style = 'LightShading-Accent1'
			hdr_cells = table.rows[0].cells
			hdr_cells[0].text = 'Item'
			hdr_cells[1].text = 'Price'
			sumtotal = 0
			currtoollist = self.estimatejobsDict["toollist_"+str(i)]
			currpricelist = self.estimatejobsDict["pricelist_"+str(i)]
			
			for item in range(len(currtoollist)):
			    row_cells = table.add_row().cells
			    row_cells[0].text = str(currtoollist[item])
			    row_cells[1].text = str(currpricelist[item])
			    sumtotal += int(currpricelist[item])
			row_cells_total = table.add_row().cells
			row_cells_total[0].text = str("TOTAL COST OF ESTIMATE:    ")
			row_cells_total[1].text = str(sumtotal)
			
		file = str(name) + '.docx'
		document.save(file)

class NamePopup(BoxLayout, Button, Screen):
	filename_text_input = ObjectProperty()
	
	def isFileOriginal(self):
		inputstr = str(self.filename_text_input.text) + '.*'
		if glob.glob(inputstr):
			popup = Popup(title='File already exists!', content=Label(text='We see that there is already a file with that name. \nPlease enter a different file name!'),
				size_hint=(None, None), size=(400, 400))
			popup.open()
			bool = False
		else: bool = True
		return bool

presentation = Builder.load_file("tooldb.kv")

class NamePopup(App):
	filename = StringProperty('')
	def build(self):
		return presentation
class ToolDBApp(App):
	def build(self):
		return presentation
class EmailApp(App):
	def build(self):
		return EmailInterface()

	def on_pause(self):
		return True

NamePopup().run()
ToolDBApp().run()
EmailApp().run()
GoBack().run()